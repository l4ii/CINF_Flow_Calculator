# -*- coding: utf-8 -*-
"""
计算书导出：
公式与含 LaTeX 的说明在文档中转为易读 Unicode/普通文本，可在 Word 内用公式编辑器再排版。
"""
from datetime import datetime
import json
import math
import os
import re
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # type: ignore
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from export_markdown import program_formula_to_latex

# 公式 id → 侧栏「四大类」之一（与前端分组一致，便于导出叙述）
FORMULA_CATEGORY_TITLE = {
    "liu_dezhong": "临界流速计算",
    "wasp": "临界流速计算",
    "fei_xiangjun": "临界流速计算",
    "kronodze_pressure": "临界流速计算",
    "clear_water_friction_loss": "摩阻损失",
    "slurry_friction_loss": "摩阻损失",
    "slurry_friction_workflow": "摩阻损失",
    "density_mixing": "摩阻损失",
    "friction_loss": "摩阻损失",
    "darcy_friction": "摩阻损失",
    "darcy_friction_step1_rho1": "摩阻损失",
    "darcy_friction_step2_re": "摩阻损失",
    "darcy_friction_step3_lambda": "摩阻损失",
    "slurry_total_head": "压力与扬程",
    "clear_water_total_head": "压力与扬程",
    "centrifugal_pump_total_head": "压力与扬程",
    "positive_displacement_pump_outlet_pressure": "压力与扬程",
    "slurry_accel_energy": "加速流与消能",
    "slurry_dissipation": "加速流与消能",
    "slurry_energy_dissipation": "加速流与消能",
    "slurry_dissipation_orifice": "加速流与消能",
}

CATEGORY_INTRO = {
    "临界流速计算": (
        "用于判定或核算浆体在管道中维持悬浮输送所需的临界流速，涵盖国内常用经验公式及 B.C.克诺罗兹法等。"
        "计算结果服务于管径与流速选取、流态判别及与后续摩阻、扬程模块的衔接。"
    ),
    "摩阻损失": (
        "用于清水或浆体管道沿程水头损失、浆体密度与达西摩阻系数等配套计算，"
        "输出单位管长水力坡降或压力梯度等，可与总扬程、加压泵站设计条件对接。"
    ),
    "压力与扬程": (
        "在几何扬程、沿程与局部损失及泵站附加损失等条件下，估算浆体或清水管道所需输送压力（表压），"
        "并可生成水力坡度线示意（纵轴水头、横轴管长；安装包内置 matplotlib 导出）。"
    ),
    "加速流与消能": (
        "提供浆体加速流判别、消能水头与孔板消能等专项核算，用于运行工况校核与消能设施参数初算；"
        "具体边界条件与系数应结合项目规范与试验资料复核。"
    ),
}


def _matplotlib_stack_for_export():
    """非交互后端 (Agg) 下的 pyplot 与 rcParams；缺失或损坏时抛出 ImportError。"""
    import matplotlib as _mpl  # pyright: ignore[reportMissingImports]
    _mpl.use("Agg")
    import matplotlib.pyplot as _plt  # pyright: ignore[reportMissingImports]
    return _plt, _mpl.rcParams


# 与前端 MainContent 水力坡度图一致：密集折线点、横轴 10 等分刻度
HYDRAULIC_GRADE_CURVE_POINTS = 240
HYDRAULIC_GRADE_TICK_DIVISIONS = 10
SLURRY_HYDRAULIC_LINE = "#F59E0B"
CLEAR_HYDRAULIC_LINE = "#3B82F6"


def _hydraulic_grade_xy(l_max, h_user, loss_head_m_fn, n_points):
    """由损失水头函数得到水力坡度折线 (L, 水头 m)，与界面公式一致。"""
    total_loss = loss_head_m_fn(l_max)
    xs, ys = [], []
    for i in range(n_points + 1):
        L = l_max * i / n_points
        lh = loss_head_m_fn(L)
        xs.append(L)
        ys.append(h_user + total_loss - lh)
    return xs, ys


class WordExporter:
    """Word 文档导出器"""
    
    # 类变量：存储每天的导出次数 {日期字符串: 次数}
    _daily_export_count = {}
    _current_date = None
    
    def __init__(self):
        # 获取当前文件所在目录（backend目录）
        current_dir = os.path.dirname(os.path.abspath(__file__))
        # 获取项目根目录（backend的父目录）
        project_root = os.path.dirname(current_dir)
        # exports目录放在项目根目录
        self.output_dir = os.path.join(project_root, "exports")
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _tex_ish_to_plain(self, s: str) -> str:
        """将 API 中带 $…$、\\rho 等的字符串转为 Word 中易读的 Unicode 文本（非 OMML 公式对象）。"""
        if s is None:
            return ""
        t = str(s)
        t = re.sub(r"\\text\{([^}]*)\}", r"\1", t)
        t = t.replace("\\quad", " ").replace("\\,", " ")
        seq = [
            ("\\omega_s", "ωs"),
            ("\\rho_w", "ρw"),
            ("\\rho_k", "ρk"),
            ("\\rho_g", "ρg"),
            ("\\rho_s", "ρs"),
            ("\\Delta", "Δ"),
            ("\\omega", "ω"),
            ("\\lambda", "λ"),
            ("\\varepsilon", "ε"),
            ("\\cdot", "·"),
            ("\\times", "×"),
            ("\\approx", "≈"),
            ("\\mathrm", ""),
            ("\\text", ""),
        ]
        for a, b in seq:
            t = t.replace(a, b)
        t = re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1)/(\2)", t)
        t = t.replace("$", "")
        t = re.sub(r"[{}]", "", t)
        t = t.replace("\\", "")
        return re.sub(r"\s+", " ", t).strip()

    def _get_app_title_version(self):
        title = (os.environ.get("FLOW_CALC_APP_TITLE") or "").strip()
        version = (os.environ.get("FLOW_CALC_APP_VERSION") or "").strip()
        if title and version:
            return title, version
        here = os.path.dirname(os.path.abspath(__file__))
        for rel in ("..", os.path.join("..", "..")):
            pkg_path = os.path.normpath(os.path.join(here, rel, "package.json"))
            if not os.path.isfile(pkg_path):
                continue
            try:
                with open(pkg_path, encoding="utf-8") as f:
                    data = json.load(f)
                t = (data.get("description") or data.get("name") or "").strip()
                v = (data.get("version") or "").strip()
                if t and v:
                    return t, v
                if v:
                    return t or "CINF长沙院浆体计算软件", v
            except (OSError, json.JSONDecodeError):
                continue
        return "CINF长沙院浆体计算软件", version or "—"

    def _apply_docx_header_to_document(self, doc: Document) -> None:
        title, ver = self._get_app_title_version()
        line = f"{title}　　版本 {ver}"
        for section in doc.sections:
            hdr = section.header
            hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            hp.text = line
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hp.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
                try:
                    run.font._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                except (AttributeError, TypeError):
                    pass

    def _get_export_count(self):
        """获取当天的导出次数并递增"""
        today = datetime.now().strftime("%Y%m%d")
        
        # 如果是新的一天，重置计数器
        if WordExporter._current_date != today:
            WordExporter._current_date = today
            WordExporter._daily_export_count[today] = 0
        
        # 递增计数器
        WordExporter._daily_export_count[today] += 1
        return WordExporter._daily_export_count[today]
    
    def export(
        self,
        formula_id,
        formula_info,
        parameters,
        result,
        save_path=None,
    ):
        """导出计算书到Word文档。save_path 为 None 时保存到 exports 目录；否则保存到用户指定路径。"""
        try:
            doc = Document()
            
            # 正文：小四（12 pt）；西文 Times New Roman，中文仿宋
            self._setup_document_style(doc)
            
            module_name = formula_info.get("name", "计算模块")
            title = doc.add_heading(f"浆体管道水力计算书 · {module_name}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.size = Pt(18)
                run.bold = True
                self._set_font(run)

            # 一、软件简要说明
            self._add_software_intro(doc)
            self._insert_page_break(doc)

            # 二、本次计算概况（功能大类 + 模块 + 时间）
            self._add_session_overview(doc, formula_id, formula_info)
            self._insert_page_break(doc)

            # 三、计算公式与模块说明
            self._add_formula_section(doc, formula_info)
            self._insert_page_break(doc)

            # 四、输入参数（含符号、取值、单位、说明）
            self._add_parameters_section(doc, parameters, formula_info)
            self._insert_page_break(doc)

            # 五、中间计算量（有则输出）
            self._add_intermediate_results(doc, result)
            if result.get("intermediate"):
                self._insert_page_break(doc)

            # 六、计算成果
            self._add_result_section(doc, result, formula_id)
            self._insert_page_break(doc)

            # 七、计算过程推演
            self._add_calculation_process(
                doc,
                formula_id,
                formula_info,
                parameters,
                result,
            )

            # 八、附录：软件功能总览（四大类）
            self._add_software_promotion(doc)

            self._apply_docx_header_to_document(doc)

            if save_path:
                # 用户指定路径（另存为）：直接保存到该路径
                save_path = os.path.abspath(save_path)
                parent = os.path.dirname(save_path)
                if parent and not os.path.exists(parent):
                    os.makedirs(parent, exist_ok=True)
                doc.save(save_path)
                return save_path
            
            # 保存到 exports 目录（兼容旧逻辑）
            timestamp = datetime.now().strftime("%Y%m%d")
            formula_name = formula_info.get('name', 'unknown').replace(' ', '').replace('/', '_')
            export_count = self._get_export_count()
            filename = f"长沙院浆体计算_{formula_name}_{timestamp}_{export_count:03d}.docx"
            file_path = os.path.join(self.output_dir, filename)
            
            # 如果文件已存在，尝试删除或重命名
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except PermissionError:
                    # 如果无法删除（可能被打开），尝试使用带时间戳的文件名
                    import time
                    timestamp_ms = int(time.time() * 1000) % 10000
                    filename = f"长沙院浆体计算_{formula_name}_{timestamp}_{export_count:03d}_{timestamp_ms}.docx"
                    file_path = os.path.join(self.output_dir, filename)
            
            # 确保目录存在且有写权限
            if not os.path.exists(self.output_dir):
                os.makedirs(self.output_dir, exist_ok=True)
            
            # 保存文件，如果失败则重试
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    doc.save(file_path)
                    break
                except PermissionError as e:
                    if attempt < max_retries - 1:
                        # 如果文件被占用，尝试使用不同的文件名
                        import time
                        timestamp_ms = int(time.time() * 1000) % 10000
                        filename = f"长沙院浆体计算_{formula_name}_{timestamp}_{export_count:03d}_{timestamp_ms}.docx"
                        file_path = os.path.join(self.output_dir, filename)
                        time.sleep(0.5)  # 等待0.5秒后重试
                    else:
                        raise Exception(f"无法保存文件，可能文件正在被其他程序打开: {file_path}")
            
            return file_path
        except Exception as e:
            import traceback
            error_msg = f"导出Word文档时出错: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)
            raise Exception(f"导出失败: {str(e)}")
    
    def _setup_document_style(self, doc):
        """正文默认：小四（12 pt），中文仿宋、西文 Times New Roman"""
        style = doc.styles['Normal']
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        try:
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "仿宋")
        except (AttributeError, TypeError):
            pass

    def _insert_page_break(self, doc):
        """段后分页，便于分章阅读"""
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_break(WD_BREAK.PAGE)

    def _set_font(self, run, chinese_font='仿宋', english_font='Times New Roman'):
        """设置run的字体：中文用指定中文字体，英文用指定英文字体"""
        run.font.name = english_font
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)
        except (AttributeError, TypeError):
            pass

    def _category_title_for(self, formula_id: str) -> str:
        return FORMULA_CATEGORY_TITLE.get(formula_id, "其他功能")

    def _style_section_heading(self, doc, text: str):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        self._set_font(run)
    
    def _add_software_intro(self, doc):
        """一、软件简要说明（不含本次算例细节）"""
        self._style_section_heading(doc, "一、软件简要说明")
        title, ver = self._get_app_title_version()
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(f"{title}（版本 {ver}）")
        sr.bold = True
        sr.font.size = Pt(12)
        self._set_font(sr)

        intro_p = doc.add_paragraph()
        intro_run = intro_p.add_run(
            "本计算书由长沙有色冶金设计研究院有限公司研发的浆体与清水管道水力计算软件自动生成。"
            "软件面向工程设计中的水力校核与资料整理，计算结论应与现行规范、试验资料及专业判断结合使用。"
        )
        self._set_font(intro_run)
        intro_p.paragraph_format.first_line_indent = Pt(24)
        intro_p.paragraph_format.line_spacing = 1.25

        auth_p = doc.add_paragraph()
        auth_run = auth_p.add_run(
            "设计使用与授权：若将本计算书所列结果、或依本软件功能得到的参数与指标，作为工程设计依据、设备选型或对外技术条件，"
            "或用于对设计起结论性、实质性指导的，须同时具备与「长沙有色冶金设计研究院有限公司」合法有效且与项目范围、用途相符的正式合同"
            "（如技术服务、工程咨询、设计等）或该院就具体项目出具的书面授权。本计算书及软件使用许可不替代、不构成上述合同或院方对具体项目的认可。"
            "未经该院书面同意，不得以长沙有色院或本公司名义将本计算书内容用于正式报审、对外技术承诺或具有担保性质的表述。"
        )
        self._set_font(auth_run)
        auth_p.paragraph_format.first_line_indent = Pt(24)
        auth_p.paragraph_format.line_spacing = 1.25

        scope_p = doc.add_paragraph()
        scope_run = scope_p.add_run("平台功能按业务板块划分为四类：")
        scope_run.bold = True
        self._set_font(scope_run)

        blocks = [
            "（1）临界流速计算：多种经验公式与克诺罗兹法等，用于流速与管径相关判别；",
            "（2）摩阻损失：清水海澄–威廉式、浆体达西–魏斯巴赫型水力坡降及密度混合、摩阻系数等配套计算；",
            "（3）压力与扬程：浆体/清水总扬程（输送压力）及压力—管长特性曲线示意；",
            "（4）加速流与消能：加速流判据、消能水头与孔板消能等专项核算。",
        ]
        for line in blocks:
            lp = doc.add_paragraph(line)
            for run in lp.runs:
                self._set_font(run)
            lp.paragraph_format.first_line_indent = Pt(24)
            lp.paragraph_format.line_spacing = 1.25

        tail = doc.add_paragraph()
        tail_run = tail.add_run(
            "本文件后续章节给出本次所选模块的公式、输入、中间量、成果及推演过程；"
            "附录中对上述四类功能作汇总说明。公式与符号在 Word 中以 Unicode 及普通文本为主，必要时可在 Word 中改用公式对象排版。"
        )
        self._set_font(tail_run)
        tail.paragraph_format.first_line_indent = Pt(24)
        tail.paragraph_format.line_spacing = 1.25

    def _add_session_overview(self, doc, formula_id, formula_info):
        """二、本次计算概况：功能大类 + 模块名称 + 大类说明"""
        self._style_section_heading(doc, "二、本次计算概况")
        cat = self._category_title_for(formula_id)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Light Grid Accent 1"
        headers = [("项目", "内容"), ("功能大类", cat), ("计算模块", formula_info.get("name", "—")), ("导出时间", datetime.now().strftime("%Y年%m月%d日 %H:%M:%S"))]
        for i, (a, b) in enumerate(headers):
            info_table.cell(i, 0).text = a
            info_table.cell(i, 1).text = b
            for cell in info_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if i == 0:
                            run.bold = True
                        self._set_font(run)

        doc.add_paragraph()
        intro_text = CATEGORY_INTRO.get(cat)
        if intro_text:
            cp = doc.add_paragraph(self._tex_ish_to_plain(intro_text))
            for run in cp.runs:
                self._set_font(run)
            cp.paragraph_format.first_line_indent = Pt(24)
            cp.paragraph_format.line_spacing = 1.25

        desc = formula_info.get("description")
        if desc and str(desc).strip():
            doc.add_paragraph()
            hp = doc.add_paragraph()
            hr = hp.add_run("本模块要点（摘要）")
            hr.bold = True
            self._set_font(hr)
            excerpt = str(desc).strip()
            if len(excerpt) > 1200:
                excerpt = excerpt[:1200] + "……（以下略，全文见「计算公式与模块说明」节）"
            for block in excerpt.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                dp = doc.add_paragraph(self._tex_ish_to_plain(block))
                for run in dp.runs:
                    self._set_font(run)
                dp.paragraph_format.first_line_indent = Pt(24)
                dp.paragraph_format.line_spacing = 1.25
    
    def _add_formula_section(self, doc, formula_info):
        """三、计算公式（程序式 → 可读文本）与模块完整说明"""
        self._style_section_heading(doc, "三、计算公式与模块说明")

        formula_name_p = doc.add_paragraph()
        formula_name_run = formula_name_p.add_run(f"模块名称：{formula_info.get('name', '未知')}")
        formula_name_run.bold = True
        self._set_font(formula_name_run)

        doc.add_paragraph()
        raw = formula_info.get("formula") or ""
        latex_frag = (formula_info.get("formula_latex") or "").strip() or program_formula_to_latex(raw)
        formula_display = self._tex_ish_to_plain(f"${latex_frag}$") if latex_frag else str(raw)

        formula_p = doc.add_paragraph()
        formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formula_run = formula_p.add_run(formula_display)
        formula_run.font.size = Pt(14)
        formula_run.font.name = "Times New Roman"
        self._set_font(formula_run)

        desc = formula_info.get("description")
        if desc:
            doc.add_paragraph()
            for block in str(desc).split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                desc_p = doc.add_paragraph(self._tex_ish_to_plain(block))
                for run in desc_p.runs:
                    self._set_font(run)
                desc_p.paragraph_format.first_line_indent = Pt(24)
                desc_p.paragraph_format.line_spacing = 1.25
    
    def _insert_math_formula(self, paragraph, formula):
        """使用OMML格式插入Word数学公式"""
        # 将公式转换为OMML格式
        omml_xml = self._convert_to_omml(formula)
        try:
            omml_element = parse_xml(omml_xml)
            paragraph._p.append(omml_element)
        except Exception as e:
            # 如果OMML插入失败，回退到文本格式
            print(f"插入数学公式失败，使用文本格式: {e}")
            formula_run = paragraph.add_run(formula)
            formula_run.font.size = Pt(14)
            formula_run.font.name = 'Cambria Math'
            self._set_font(formula_run)
    
    def _convert_to_omml(self, formula):
        """将公式字符串转换为OMML XML格式"""
        # 解析公式并转换为OMML格式
        # 这是一个简化的转换，可以根据样本文档进一步优化
        
        # 处理变量下标：Vc -> V_c, Cv -> C_V（体积浓度下标大写 V）, d85 -> d_85, d90 -> d_90
        formula = re.sub(r'Vc', 'V_c', formula)
        formula = re.sub(r'Cv', 'C_V', formula)
        formula = re.sub(r'd85', 'd_85', formula)
        formula = re.sub(r'd90', 'd_90', formula)
        formula = re.sub(r'ω_s', 'ω_s', formula)
        formula = re.sub(r'rho_g', 'ρ_g', formula)
        formula = re.sub(r'rho_k', 'ρ_k', formula)
        
        # 构建OMML XML
        # 这里先创建一个基本的OMML结构，具体格式可以根据样本文档调整
        omml_parts = []
        
        # 分割公式为各个部分
        parts = re.split(r'(\s*=\s*|\s*\*\s*|\s*\[\s*|\s*\]\s*|\s*\(\s*|\s*\)\s*|\s*\^\s*|\s*/\s*)', formula)
        
        omml_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        
        # 简化处理：先使用文本格式，等样本文档后再优化
        # 将公式转换为Unicode数学符号格式
        formula_display = formula.replace('*', '·').replace('^', '^')
        
        # 创建文本run
        omml_xml += '<m:r><m:t xml:space="preserve">' + self._escape_xml(formula_display) + '</m:t></m:r>'
        
        omml_xml += '</m:oMath>'
        
        return omml_xml
    
    def _escape_xml(self, text):
        """转义XML特殊字符"""
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&apos;'))
    
    def _add_parameters_section(self, doc, parameters, formula_info):
        """四、输入参数：符号（含中文说明）、取值、单位、字段说明"""
        self._style_section_heading(doc, "四、输入参数与取值")

        formula_params = formula_info.get("parameters", [])
        raw = parameters or {}
        valid_params = {}
        for k, v in raw.items():
            if v is None or v == "":
                continue
            if isinstance(v, float) and math.isnan(v):
                continue
            valid_params[k] = v

        def _fmt_val(value):
            if isinstance(value, (int, float)):
                return f"{value:.6f}".rstrip("0").rstrip(".")
            return str(value)

        param_def_by_name = {p.get("name"): p for p in formula_params if p.get("name")}

        formula_param_names = [p.get("name") for p in formula_params if p.get("name")]
        ordered_names = [n for n in formula_param_names if n in valid_params]
        extra_keys = [k for k in valid_params if k not in set(formula_param_names)]
        extra_keys.sort()
        all_names = ordered_names + extra_keys
        total_rows = len(all_names)

        if total_rows == 0:
            no_params_p = doc.add_paragraph("（本模块无独立数值参数或由界面分步汇总，详见计算过程节。）")
            for run in no_params_p.runs:
                self._set_font(run)
            return

        param_table = doc.add_table(rows=total_rows + 1, cols=4)
        param_table.style = "Light Grid Accent 1"

        header_cells = param_table.rows[0].cells
        header_cells[0].text = "参数（符号与含义）"
        header_cells[1].text = "取值"
        header_cells[2].text = "单位"
        header_cells[3].text = "说明"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    self._set_font(run)

        row = 1
        for name in all_names:
            value = valid_params[name]
            pdef = param_def_by_name.get(name)
            label = self._tex_ish_to_plain(pdef.get("label", name)) if pdef else name
            unit = (pdef.get("unit") if pdef else None) or self._get_unit(name)
            desc = (pdef.get("description") if pdef else None) or "—"
            if desc and str(desc).strip():
                desc = self._tex_ish_to_plain(str(desc).strip())
            else:
                desc = "—"

            param_table.cell(row, 0).text = label
            param_table.cell(row, 1).text = _fmt_val(value)
            param_table.cell(row, 2).text = unit or "—"
            param_table.cell(row, 3).text = desc
            for cell in param_table.rows[row].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._set_font(run)
            row += 1
    
    def _add_intermediate_results(self, doc, result):
        """五、中间计算量"""
        intermediate = result.get('intermediate', {})
        if not intermediate:
            return

        self._style_section_heading(doc, "五、中间计算量")
        
        # 创建中间结果表格
        intermediate_table = doc.add_table(rows=len(intermediate) + 1, cols=2)
        intermediate_table.style = 'Light Grid Accent 1'
        
        # 表头
        header_cells = intermediate_table.rows[0].cells
        header_cells[0].text = "项目"
        header_cells[1].text = "数值"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    self._set_font(run)
        
        # 填充中间结果
        row = 1
        for key, value in intermediate.items():
            # 格式化键名（转换为中文标签）
            label = self._get_intermediate_label(key)
            intermediate_table.cell(row, 0).text = label
            
            # 格式化数值
            if isinstance(value, (int, float)):
                if abs(value) < 0.001:
                    value_str = f"{value:.6e}"
                elif abs(value) < 1:
                    value_str = f"{value:.6f}".rstrip('0').rstrip('.')
                else:
                    value_str = f"{value:.4f}".rstrip('0').rstrip('.')
            else:
                value_str = str(value)
            
            intermediate_table.cell(row, 1).text = value_str
            # 设置该行所有单元格的字体
            for cell in intermediate_table.rows[row].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._set_font(run)
            row += 1
    
    def _get_intermediate_label(self, key):
        """获取中间计算项的中文标签"""
        labels = {
            'delta_rho_ratio': '相对密度差 Δρ/ρ',
            'density_ratio': '密度比 (ps-pl)/pl',
            'core_term': '核心项 [g·D·(Δρ/ρ)·ω]^(1/3)',
            'concentration_term': '浓度修正项 C_V^(1/6)',
            'velocity_ratio_term': '速度比修正项 (ω_s/ω)^(1/6)',
            'bracket_term': '核心项 [2·g·D·(Δρ/ρ)]^(1/2)',
            'size_ratio_term': '粒径比修正项 (d85/D)^(1/6)',
            'conc_term': '浓度修正项 C_V^0.25',
            'size_term': '粒径比修正项 (d90/D)^(1/3)',
            'leading_coef': '核心系数 2.26/√λ',
            'sqrt_term': '平方根项',
            'sin_theta': 'sin(θ)',
            'coefficient': '经验系数',
            'coefficient_9_5': '经验系数 9.5',
            'coefficient_3_113': '经验系数 3.113',
            'coefficient_2_26': '经验系数 2.26',
            'g': '重力加速度 g',
            # 沿程摩阻损失（4.3.1-1）
            'numerator': '流速平方与浆体密度项 V²·ρ_k',
            'denominator': '重力与管径项 2gD·ρ_s',
            # 密度混合公式（4.3.1-2）
            'denom': '浓度与密度加权倒数项 C_w/ρ_g+(1-C_w)/ρ_s',
            # B.C.克诺罗兹法
            'step_A_Qk': '步骤A 矿浆流量 Qk',
            'step_B_DL_mm': '步骤B 临界管径 DL (mm)',
            'Cd': '重量砂水比 Cd',
            'step_C_V_L': '步骤C 临界流速 V_L',
            # 达西摩阻系数公式
            'Re': '雷诺数 Re',
            'flow_regime': '流态',
            'eps_D': '相对粗糙度 ε/D',
            # 浆体加速流及消能
            'head_diff': '左侧总水头差 (Z₁+H₁)-(Z₂+H₂)',
            'friction_loss_total': '右侧摩阻损失 iL',
            "step_1_kql": "步骤1 流量消能系数 K_QL",
            "step_2_delta_h": "步骤2 消能水头 Δh",
            "kql_numerator": "K_QL 分子项 (6.3755×10⁻⁹)·λ_d·L_s",
            "kql_denominator_d5": "K_QL 分母项 d⁵",
            "Q_squared": "Q²",
            "gravity_pressure": "重力势能压力项",
            "friction_pressure": "沿程摩擦压力项",
            "orifice_step": "孔板计算子步序号",
            "orifice_numer": "K_Qk 分子项 (1-β²)(1.142-β²)",
            "orifice_beta": "步骤1 孔径比 β",
            "orifice_K_Qk_step2": "步骤2 孔板流量消能系数 K_Qk",
            "term_0p25_Cw": "项 0.25·C_w",
            "Sigma_H_s": "装置所需压力累计 ΣH_s",
            "K_p_K_m": "分母 K_p·K_m",
            "K_p": "扬程降低率 K_p",
            "K_m": "磨蚀后扬程折损率 K_m",
            "C_w": "浆体重量浓度 C_w",
            "K_f": "压力富余系数 K_f",
            "P_k": "浆体管道输送压力 P_k（输入）",
        }
        return labels.get(key, key)
    
    def _add_result_section(self, doc, result, formula_id=None):
        """六、计算成果。根据 formula_id 显示 Vc、i_k、rho_k 等"""
        self._style_section_heading(doc, "六、计算成果")
        
        result_table = doc.add_table(rows=3, cols=2)
        result_table.style = 'Light Grid Accent 1'
        
        # 设置表头样式
        header_cells = result_table.rows[0].cells
        header_cells[0].text = '项目'
        header_cells[1].text = '结果'
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    self._set_font(run)
        
        # 根据公式类型显示对应结果
        if formula_id == 'friction_loss':
            item_label = '沿程摩阻损失 i_k'
            value = result.get('i_k', 'N/A')
        elif formula_id == 'density_mixing':
            item_label = '浆体密度 ρ_k'
            value = result.get('rho_k', 'N/A')
        elif formula_id == 'darcy_friction':
            item_label = '达西摩阻系数'
            rho_1 = result.get('rho_1', 'N/A')
            Re_B = result.get('Re_B', 'N/A')
            lam = result.get('lambda_coef', 'N/A')
            flow_regime = result.get('intermediate', {}).get('flow_regime', '')
            value = f"ρ₁={rho_1} t/m³，ReB={Re_B}，λ={lam}" + (f"（{flow_regime}）" if flow_regime else "")
        elif formula_id == "darcy_friction_step1_rho1":
            item_label = "混合物密度 ρ₁"
            value = f"{result.get('rho_1', 'N/A')} t/m³"
        elif formula_id == "darcy_friction_step2_re":
            item_label = "混合物雷诺数 Re_B"
            value = str(result.get("Re_B", "N/A"))
        elif formula_id == "darcy_friction_step3_lambda":
            item_label = "达西摩阻系数 λ"
            value = str(result.get("lambda_coef", "N/A"))
        elif formula_id == 'slurry_accel_energy':
            item_label = '浆体加速流条件'
            value = '满足' if result.get('condition_met') else '不满足'
        elif formula_id in ('slurry_dissipation', 'slurry_energy_dissipation'):
            item_label = '浆体消能水头 Δh'
            delta_h = result.get('delta_h', result.get('intermediate', {}).get('step_2_delta_h', 'N/A'))
            kql = result.get('K_QL', result.get('intermediate', {}).get('step_1_kql', 'N/A'))
            value = f"Δh = {delta_h} m，K_QL = {kql} h²/m⁵"
        elif formula_id == 'slurry_dissipation_orifice':
            item_label = '孔板消能水头 Δh'
            delta_h = result.get('delta_h', 'N/A')
            value = f"Δh = {delta_h} m（K_Qk、Q 见「输入参数与取值」节）"
        elif formula_id == 'slurry_friction_loss':
            item_label = '浆体摩阻损失'
            rho_k = result.get('rho_k', 'N/A')
            i_k = result.get('i_k', 'N/A')
            value = f"ρ_k = {rho_k} t/m³，i_k = {i_k} mH₂O/m"
        elif formula_id == 'slurry_friction_workflow':
            item_label = '浆体摩阻损失（分步）'
            rho_k = result.get('rho_k', 'N/A')
            i_k = result.get('i_k', 'N/A')
            value = f"ρ_k = {rho_k} t/m³，i_k = {i_k} mH₂O/m"
        elif formula_id == 'slurry_total_head':
            item_label = '浆体管道输送压力 Pk'
            ht = result.get('H_total', 'N/A')
            value = f"{ht} kPa" if ht != 'N/A' else 'N/A'
        elif formula_id == 'clear_water_total_head':
            item_label = '清水管道输送压力 Pw'
            ht = result.get('H_total', 'N/A')
            value = f"{ht} kPa" if ht != 'N/A' else 'N/A'
        elif formula_id == 'centrifugal_pump_total_head':
            item_label = '主泵扬送清水的总扬程 H_b（液柱）'
            ht = result.get('H_total', 'N/A')
            value = f"{ht} m" if ht != 'N/A' else 'N/A'
        elif formula_id == 'positive_displacement_pump_outlet_pressure':
            item_label = '容积式泵总扬程 P_b（压力）'
            pb = result.get('P_b', result.get('H_total', 'N/A'))
            value = f"{pb} kPa" if pb != 'N/A' else 'N/A'
        elif formula_id == 'clear_water_friction_loss':
            item_label = '单位长度水头损失 i'
            i_val = result.get('i', 'N/A')
            if isinstance(i_val, (int, float)):
                value = f"{i_val} kPa/m"
            else:
                value = str(i_val)
        else:
            item_label = '临界流速 Vc'
            value = result.get('Vc', 'N/A')
        
        if isinstance(value, (int, float)):
            value_display = f"{value:.4f}".rstrip('0').rstrip('.')
        else:
            value_display = str(value)
        
        result_table.cell(1, 0).text = item_label
        unit_suffix = result.get("unit", "")
        embed_unit_ids = (
            "slurry_friction_loss",
            "slurry_friction_workflow",
            "darcy_friction",
            "darcy_friction_step1_rho1",
            "darcy_friction_step2_re",
            "darcy_friction_step3_lambda",
            "slurry_dissipation",
            "slurry_energy_dissipation",
            "slurry_dissipation_orifice",
            "slurry_total_head",
            "clear_water_total_head",
            "centrifugal_pump_total_head",
            "positive_displacement_pump_outlet_pressure",
            "clear_water_friction_loss",
        )
        if formula_id in embed_unit_ids:
            result_table.cell(1, 1).text = value_display
        else:
            result_table.cell(1, 1).text = f"{value_display} {unit_suffix}".strip() if unit_suffix else value_display
        # 设置该行所有单元格的字体
        for cell in result_table.rows[1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._set_font(run)
        
        # 备注
        result_table.cell(2, 0).text = "备注"
        result_table.cell(2, 1).text = (
            "本表数值由当前输入与所选公式按程序逻辑计算得出，用于设计辅助与资料整理；"
            "实施阶段请结合现场条件、规范及专业判断复核。"
            "若将本表结果或依本软件功能形成的指标作为工程设计或对外正式依据，须同时具备与「长沙有色冶金设计研究院有限公司」"
            "合法有效且与项目相符的正式合同或书面项目授权；本计算书不替代该等合同或授权。"
        )
        # 设置该行所有单元格的字体
        for cell in result_table.rows[2].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._set_font(run)
    
    def _add_calculation_process(
        self,
        doc,
        formula_id,
        formula_info,
        parameters,
        result,
    ):
        """七、计算过程推演"""
        self._style_section_heading(doc, "七、计算过程推演")

        intro = doc.add_paragraph(
            "以下按本模块常用书写顺序列出主要代入关系与中间量；若上一节「中间计算量」有表，则本节数值与其一致，便于对照复核。"
        )
        for run in intro.runs:
            self._set_font(run)
        intro.paragraph_format.first_line_indent = Pt(24)
        intro.paragraph_format.line_spacing = 1.25

        # 根据公式ID添加详细计算步骤
        if formula_id == "liu_dezhong":
            self._add_liu_dezhong_process(doc, parameters, result)
        elif formula_id == "wasp":
            self._add_wasp_process(doc, parameters, result)
        elif formula_id == "fei_xiangjun":
            self._add_fei_xiangjun_process(doc, parameters, result)
        elif formula_id == "kronodze_pressure":
            self._add_kronodze_pressure_process(doc, parameters, result)
        elif formula_id == "friction_loss":
            self._add_friction_loss_process(doc, parameters, result)
        elif formula_id == "density_mixing":
            self._add_density_mixing_process(doc, parameters, result)
        elif formula_id == "darcy_friction":
            self._add_darcy_friction_process(doc, parameters, result)
        elif formula_id in ("darcy_friction_step1_rho1", "darcy_friction_step2_re", "darcy_friction_step3_lambda"):
            self._add_darcy_substep_process(doc, formula_id, parameters, result)
        elif formula_id == "slurry_accel_energy":
            self._add_slurry_accel_energy_process(doc, parameters, result)
        elif formula_id in ("slurry_dissipation", "slurry_energy_dissipation"):
            self._add_slurry_dissipation_process(doc, parameters, result)
        elif formula_id == "slurry_dissipation_orifice":
            self._add_slurry_dissipation_orifice_process(doc, parameters, result)
        elif formula_id == "slurry_friction_loss":
            self._add_slurry_friction_loss_process(doc, parameters, result)
        elif formula_id == "slurry_friction_workflow":
            intro_wf = doc.add_paragraph(
                "本模块在界面内按步骤完成：浆体密度 ρ_k、混合物密度 ρ₁、混合物雷诺数 Re_B、达西摩阻系数 λ 与单位管长水力坡降 i_k；"
                "下列推演对应最终合并代入关系，与界面分步结果一致。"
            )
            for run in intro_wf.runs:
                self._set_font(run)
            intro_wf.paragraph_format.first_line_indent = Pt(24)
            self._add_slurry_friction_loss_process(doc, parameters, result)
        elif formula_id == "slurry_total_head":
            self._add_slurry_total_head_process(doc, parameters, result)
        elif formula_id == "clear_water_total_head":
            self._add_clear_water_total_head_process(doc, parameters, result)
        elif formula_id == "centrifugal_pump_total_head":
            self._add_centrifugal_pump_total_head_process(doc, parameters, result)
        elif formula_id == "positive_displacement_pump_outlet_pressure":
            self._add_positive_displacement_pump_process(doc, parameters, result)
        elif formula_id == "clear_water_friction_loss":
            self._add_clear_water_friction_loss_process(doc, parameters, result)
    
    def _add_clear_water_friction_loss_process(self, doc, parameters, result):
        """清水海澄–威廉沿程摩阻：书写为 105·C_h^(-1.85)·…，计算用参数 K_hw（默认 105）"""
        intermediate = result.get("intermediate", {}) or {}
        ch = parameters.get("C_h", "N/A")
        k_hw = parameters.get("K_hw", 105)
        dj = parameters.get("d_j", "N/A")
        qg = parameters.get("q_g", "N/A")
        process_texts = [
            f"1. 输入：K_hw = {k_hw}，C_h = {ch}，d_j = {dj} m，q_g = {qg} m³/s",
            f"2. C_h^(-1.85) = {intermediate.get('clear_hw_ch_pow', 'N/A')}",
            f"3. d_j^(-4.87) = {intermediate.get('clear_hw_dj_pow', 'N/A')}",
            f"4. q_g^1.85 = {intermediate.get('clear_hw_qg_pow', 'N/A')}",
            f"5. i = {k_hw} × 以上三项之积（式中常数 105 对应参数 K_hw）= {result.get('i', 'N/A')} kPa/m",
        ]
        for text in process_texts:
            para = doc.add_paragraph(text)
            for run in para.runs:
                self._set_font(run)

    def _add_liu_dezhong_process(self, doc, parameters, result):
        """添加刘德忠公式计算过程"""
        intermediate = result.get('intermediate', {})
        rho_g = parameters.get('rho_g', 'N/A')
        rho_k = parameters.get('rho_k', 'N/A')
        coefficient = intermediate.get('coefficient', parameters.get('coefficient_9_5', 9.5))
        
        process_texts = [
            f"1. 计算相对密度差: Δρ/ρ = ({rho_g} - {rho_k})/{rho_k} = {intermediate.get('delta_rho_ratio', 'N/A')}",
            f"2. 计算核心项: [g·D·(Δρ/ρ)·ω]^(1/3) = {intermediate.get('core_term', 'N/A')}",
            f"3. 计算浓度修正项: C_V^(1/6) = {intermediate.get('concentration_term', 'N/A')}",
            f"4. 计算速度比修正项: (ω_s/ω)^(1/6) = {intermediate.get('velocity_ratio_term', 'N/A')}",
            f"5. 计算临界流速: Vc = {coefficient} × {intermediate.get('core_term', 'N/A')} × {intermediate.get('concentration_term', 'N/A')} × {intermediate.get('velocity_ratio_term', 'N/A')}",
            f"   Vc = {result.get('Vc', 'N/A')} m/s"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)
    
    def _add_wasp_process(self, doc, parameters, result):
        """添加瓦斯普公式计算过程"""
        intermediate = result.get('intermediate', {})
        coefficient = intermediate.get('coefficient', parameters.get('coefficient_3_113', 3.113))
        rho_g = parameters.get('rho_g', 'N/A')
        rho_k = parameters.get('rho_k', 'N/A')
        
        process_texts = [
            f"1. 计算相对密度差: Δρ/ρ = ({rho_g} - {rho_k})/{rho_k} = {intermediate.get('delta_rho_ratio', 'N/A')}",
            f"2. 计算核心项: [2·g·D·(Δρ/ρ)]^(1/2) = {intermediate.get('bracket_term', 'N/A')}",
            f"3. 计算浓度修正项: C_V^0.1858 = {intermediate.get('concentration_term', 'N/A')}",
            f"4. 计算粒径比修正项: (d85/D)^(1/6) = {intermediate.get('size_ratio_term', 'N/A')}",
            f"5. 计算临界流速: Vc = {coefficient} × {intermediate.get('concentration_term', 'N/A')} × {intermediate.get('bracket_term', 'N/A')} × {intermediate.get('size_ratio_term', 'N/A')}",
            f"   Vc = {result.get('Vc', 'N/A')} m/s"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)
    
    def _add_fei_xiangjun_process(self, doc, parameters, result):
        """添加费祥俊公式计算过程"""
        intermediate = result.get('intermediate', {})
        rho_g = parameters.get('rho_g', 'N/A')
        rho_k = parameters.get('rho_k', 'N/A')
        lambda_coef = parameters.get('lambda_coef', 'N/A')
        coefficient = intermediate.get('coefficient_2_26', parameters.get('coefficient_2_26', 2.26))
        
        process_texts = [
            f"1. 计算相对密度差: Δρ/ρ = ({rho_g} - {rho_k})/{rho_k} = {intermediate.get('delta_rho_ratio', 'N/A')}",
            f"2. 计算核心系数: 2.26/√λ = {coefficient}/√{lambda_coef} = {intermediate.get('leading_coef', 'N/A')}",
            f"3. 计算核心项: [g·D·(Δρ/ρ)]^(1/2) = {intermediate.get('bracket_term', 'N/A')}",
            f"4. 计算浓度修正项: C_V^0.25 = {intermediate.get('conc_term', 'N/A')}",
            f"5. 计算粒径比修正项: (d90/D)^(1/3) = {intermediate.get('size_term', 'N/A')}",
            f"6. 计算临界流速: Vc = {intermediate.get('leading_coef', 'N/A')} × {intermediate.get('bracket_term', 'N/A')} × {intermediate.get('conc_term', 'N/A')} × {intermediate.get('size_term', 'N/A')}",
            f"   Vc = {result.get('Vc', 'N/A')} m/s"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)
    
    def _add_kronodze_pressure_process(self, doc, parameters, result):
        """添加 B.C.克诺罗兹法 三步计算过程"""
        intermediate = result.get('intermediate', {})
        K = parameters.get('K', 1.1)
        G = parameters.get('G', 'N/A')
        W = parameters.get('W', 'N/A')
        rho_g = parameters.get('rho_g', 'N/A')
        dp = parameters.get('dp', 'N/A')
        beta = parameters.get('beta', 1.0)
        Qk = intermediate.get('step_A_Qk', 'N/A')
        DL = intermediate.get('step_B_DL_mm', 'N/A')
        Cd = intermediate.get('Cd', 'N/A')
        process_texts = [
            "A) 计算矿浆流量 Qk：",
            f"   Qk = K·W·(1/ρg + G/W) = {K}×{W}×(1/{rho_g} + {G}/{W}) = {Qk}",
            "B) 计算临界管径 DL（按尾矿加权平均粒径 dp 选用公式，由 Qk 反解）：",
            f"   dp = {dp} mm，重量砂水比 Cd = W/G×100 = {Cd}，得 DL = {DL} mm",
            "C) 计算临界流速 V_L：",
            f"   V_L = 0.255β(1 + 2.48·³√(Cd)·⁴√(DL)) = {result.get('Vc', 'N/A')} m/s"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)
    
    def _add_friction_loss_process(self, doc, parameters, result):
        """添加沿程摩阻损失(4.3.1-1)计算过程"""
        lambda_coef = parameters.get('lambda_coef', 'N/A')
        V = parameters.get('V', 'N/A')
        rho_k = parameters.get('rho_k', 'N/A')
        D = parameters.get('D', 'N/A')
        rho_s = parameters.get('rho_s', 'N/A')
        g = parameters.get('g', 9.81)
        i_k = result.get('i_k', 'N/A')
        process_texts = [
            f"公式(4.3.1-1): i_k = λ·(V²·ρ_k)/(2gD·ρ_s)",
            f"1. 代入: λ={lambda_coef}, V={V} m/s, ρ_k={rho_k} t/m³, D={D} m, ρ_s={rho_s} t/m³, g={g} m/s²",
            f"2. 沿程摩阻损失: i_k = {i_k} mH₂O/m"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)
    
    def _add_density_mixing_process(self, doc, parameters, result):
        """添加密度混合公式(4.3.1-2)计算过程"""
        C_w = parameters.get('C_w', 'N/A')
        rho_g = parameters.get('rho_g', 'N/A')
        rho_s = parameters.get('rho_s', 'N/A')
        rho_k = result.get('rho_k', 'N/A')
        process_texts = [
            f"公式(4.3.1-2): ρ_k = 1/(C_w/ρ_g + (1-C_w)/ρ_s)",
            f"1. 代入: C_w={C_w}, ρ_g={rho_g} t/m³, ρ_s={rho_s} t/m³",
            f"2. 浆体密度: ρ_k = {rho_k} t/m³"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

    def _add_darcy_substep_process(self, doc, formula_id, parameters, result):
        """达西链单步：ρ₁ / Re_B / λ"""
        lines = []
        if formula_id == "darcy_friction_step1_rho1":
            lines = [
                f"混合物密度：ρ₁ = {result.get('rho_1', 'N/A')} t/m³",
            ]
        elif formula_id == "darcy_friction_step2_re":
            lines = [
                f"混合物雷诺数：Re_B = {result.get('Re_B', 'N/A')}（ρ₁ = {result.get('rho_1', 'N/A')} t/m³）",
            ]
        elif formula_id == "darcy_friction_step3_lambda":
            im = result.get("intermediate") or {}
            regime = im.get("flow_regime", "")
            lines = [
                f"达西摩阻系数：λ = {result.get('lambda_coef', 'N/A')}"
                + (f"，流态：{regime}" if regime else ""),
                f"雷诺数 Re_B = {result.get('Re_B', 'N/A')}",
            ]
        for text in lines:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

    def _add_darcy_friction_process(self, doc, parameters, result):
        """添加达西摩阻系数（三步）计算过程"""
        intermediate = result.get('intermediate', {})
        rho_1 = intermediate.get('step_A_rho_1', result.get('rho_1', 'N/A'))
        Re_B = intermediate.get('step_B_Re_B', result.get('Re_B', 'N/A'))
        lam = result.get('lambda_coef', 'N/A')
        flow_regime = intermediate.get('flow_regime', 'N/A')
        rho_g, rho_s, C1v = parameters.get('rho_g'), parameters.get('rho_s'), parameters.get('C1v')
        if rho_g is not None and rho_s is not None and C1v is not None:
            step_a = [f"步骤 A: ρ₁ = ρg·C1v + (1-C1v)·ρs（t/m³）", f"代入 ρg={rho_g}, ρs={rho_s}, C1v={C1v} → ρ₁ = {rho_1} t/m³"]
        else:
            step_a = [f"步骤 A: 用户直接输入 ρ₁ = {rho_1} t/m³"]
        V, D_n, eta_1 = parameters.get('V'), parameters.get('D_n'), parameters.get('eta_1')
        if V is not None and D_n is not None and eta_1 is not None:
            step_b = ["", "步骤 B: ReB = (V·Dn·1000·ρ₁)/η₁（ρ₁ 为 t/m³）", f"代入 V={V}, Dn={D_n}, ρ₁={rho_1}, η₁={eta_1} → ReB = {Re_B}"]
        else:
            step_b = ["", f"步骤 B: 用户直接输入 ReB = {Re_B}"]
        epsilon = parameters.get('epsilon', 0.0002)
        step_c = ["", "步骤 C: 达西摩阻系数 λ", f"ReB={Re_B}，流态：{flow_regime}", f"ε={epsilon}, Dn={D_n}", f"λ = {lam}"]
        for text in step_a + step_b + step_c:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

    def _add_slurry_accel_energy_process(self, doc, parameters, result):
        """添加浆体加速流/消能计算过程"""
        intermediate = result.get('intermediate', {})
        Z1, Z2 = parameters.get('Z1', 'N/A'), parameters.get('Z2', 'N/A')
        H1, H2 = parameters.get('H1', 'N/A'), parameters.get('H2', 'N/A')
        i, L = parameters.get('i', 'N/A'), parameters.get('L', 'N/A')
        head_diff = intermediate.get('head_diff', 'N/A')
        friction_loss_total = intermediate.get('friction_loss_total', 'N/A')
        condition_met = result.get('condition_met', False)
        conclusion = '满足' if condition_met else '不满足'
        process_texts = [
            "公式(6): (Z₁ + P₁/(ρkg)) - (Z₂ + P₂/(ρkg)) > iL",
            f"1. 左侧总水头差 = (Z₁+H₁)-(Z₂+H₂) = ({Z1}+{H1})-({Z2}+{H2}) = {head_diff} m",
            f"2. 右侧摩阻损失 = i×L = {i}×{L} = {friction_loss_total} m",
            f"3. 判断: {head_diff} {'>' if condition_met else '≤'} {friction_loss_total}，条件{conclusion}"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

    def _add_slurry_dissipation_process(self, doc, parameters, result):
        """添加浆体消能两步计算过程"""
        intermediate = result.get('intermediate', {})
        lambda_d = parameters.get('lambda_d', 'N/A')
        L_s = parameters.get('L_s', 'N/A')
        d = parameters.get('d', 'N/A')
        Q = parameters.get('Q', 'N/A')
        K_QL = result.get('K_QL', intermediate.get('step_1_kql', parameters.get('K_QL', 'N/A')))
        delta_h = result.get('delta_h', intermediate.get('step_2_delta_h', 'N/A'))
        num = intermediate.get('kql_numerator')
        den = intermediate.get('kql_denominator_d5')
        Q2 = intermediate.get('Q_squared')

        process_texts = [
            "步骤1：计算沿程缩径增阻管道流量消能系数",
            "K_QL = (6.3755×10^-9)·λ_d·L_s / d^5",
        ]
        if num is not None and den is not None:
            process_texts.append(f"中间项：分子 (6.3755×10^-9)·λ_d·L_s = {num}；分母 d^5 = {den}")
        process_texts.append(f"代入 λ_d={lambda_d}, L_s={L_s}, d={d}，得到 K_QL = {K_QL} h²/m⁵")
        process_texts.extend([
            "步骤2：计算消能水头",
            "Δh = K_QL·Q²",
        ])
        if Q2 is not None:
            process_texts.append(f"中间项：Q² = {Q2}")
        process_texts.append(f"代入 K_QL={K_QL}, Q={Q}，得到 Δh = {delta_h} m")
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

    def _add_slurry_dissipation_orifice_process(self, doc, parameters, result):
        """孔板消能三步分算：按输入表中出现的量整理推演说明（允许跳过前两步仅做第 3 步）。"""
        intermediate = result.get('intermediate', {}) or {}
        d = parameters.get('d', 'N/A')
        D = parameters.get('D', 'N/A')
        beta = parameters.get('beta', intermediate.get('orifice_beta', 'N/A'))
        K_Qk_in = parameters.get('K_Qk', 'N/A')
        Q = parameters.get('Q', 'N/A')
        delta_h = result.get('delta_h', 'N/A')
        Q2 = intermediate.get('Q_squared')
        kqk2 = intermediate.get('orifice_K_Qk_step2')

        lines = [
            "孔板（节流件）局部消能：β = d/D；K_Qk = 6.3755×10^-9·(1-β²)(1.142-β²)/d⁴；Δh = K_Qk·Q²（Q 单位 m³/h，Δh 单位 m）。",
        ]
        if d != 'N/A' and D != 'N/A':
            lines.append(f"步骤1（若适用）：由 d={d} m、D={D} m 得孔径比 β = d/D = {beta}。")
        elif beta != 'N/A':
            lines.append(f"孔径比 β = {beta}（由用户直接给定或自其他途径取得）。")
        if kqk2 is not None:
            lines.append(f"步骤2（若适用）：由 β 与 d 得孔板流量消能系数 K_Qk = {kqk2} h²/m⁵。")
        lines.append(f"步骤3：代入 K_Qk = {K_Qk_in} h²/m⁵、Q = {Q} m³/h。")
        if Q2 is not None:
            lines.append(f"中间量 Q² = {Q2} (m³/h)²。")
        lines.append(f"消能水头 Δh = K_Qk·Q² = {delta_h} m。")
        for text in lines:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

    def _add_slurry_friction_loss_process(self, doc, parameters, result):
        """添加浆体摩阻损失（达西-魏斯巴赫公式）计算过程"""
        rho_k = parameters.get('rho_k', result.get('rho_k', 'N/A'))
        rho_s = parameters.get('rho_s', 'N/A')
        lambda_coef = parameters.get('lambda_coef', 'N/A')
        V = parameters.get('V', 'N/A')
        D = parameters.get('D', 'N/A')
        g = parameters.get('g', 9.81)
        i_k = result.get('i_k', 'N/A')
        process_texts = [
            "达西-魏斯巴赫公式(4.3.1-1): i_k = λ·(V²·ρ_k)/(2gD·ρ_s)",
            f"1. 代入: λ={lambda_coef}, V={V} m/s, ρ_k={rho_k} t/m³, D={D} m, ρ_s={rho_s} t/m³, g={g} m/s²",
            f"2. 沿程摩阻损失: i_k = {i_k} mH₂O/m"
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)
    
    @staticmethod
    def _try_slurry_hydraulic_grade_xy(parameters):
        """与界面浆体水力坡度线相同的 (L, 水头 m) 折线；失败返回 None。"""
        try:
            l_max = float(parameters.get('L') or 0)
            if l_max <= 0:
                return None
            rho_s = float(parameters.get('rho_s') or 0)
            rho_k = float(parameters.get('rho_k') or 0)
            g = float(parameters.get('g') or 9.81)
            i_k = float(parameters.get('i_k') or 0)
            H = float(parameters.get('H') or 0)
            P_j = float(parameters.get('P_j') or 0)
            if rho_s <= 0 or rho_k <= 0 or g <= 0:
                return None

            def loss_head_m(l):
                pk = rho_s * g * i_k * l + (P_j * (l / l_max) if l_max > 0 else 0)
                return pk / (rho_k * g)

            return _hydraulic_grade_xy(l_max, H, loss_head_m, HYDRAULIC_GRADE_CURVE_POINTS)
        except (TypeError, ValueError, ZeroDivisionError):
            return None

    @staticmethod
    def _try_slurry_page_clear_hydraulic_grade_xy(parameters):
        """浆体页清水对比线：与界面一致，取当前浆体参数的 H、L、P_j、g 与 i_k，ρ_w=1 t/m³、i_w=i_k。"""
        try:
            l_max = float(parameters.get('L') or 0)
            if l_max <= 0:
                return None
            rho_w = 1.0
            g = float(parameters.get('g') or 9.81)
            i_w = float(parameters.get('i_k') or 0)
            H = float(parameters.get('H') or 0)
            P_j = float(parameters.get('P_j') or 0)
            if g <= 0 or rho_w <= 0:
                return None

            def loss_head_m(l):
                pk = rho_w * g * i_w * l + (P_j * (l / l_max) if l_max > 0 else 0)
                return pk / (rho_w * g)

            return _hydraulic_grade_xy(l_max, H, loss_head_m, HYDRAULIC_GRADE_CURVE_POINTS)
        except (TypeError, ValueError, ZeroDivisionError):
            return None

    @staticmethod
    def _try_clear_hydraulic_grade_xy(l_max, clear_parameters):
        """与界面清水水力坡度线相同；横轴用浆体侧 l_max。"""
        try:
            if l_max <= 0 or not clear_parameters:
                return None
            rho_w = float(clear_parameters.get('rho_w') or 1)
            g = float(clear_parameters.get('g') or 9.81)
            i_w = float(clear_parameters.get('i_w') or 0)
            H = float(clear_parameters.get('H') or 0)
            P_j = float(clear_parameters.get('P_j') or 0)
            if rho_w <= 0 or g <= 0:
                return None

            def loss_head_m(l):
                pk = rho_w * g * i_w * l + (P_j * (l / l_max) if l_max > 0 else 0)
                return pk / (rho_w * g)

            return _hydraulic_grade_xy(l_max, H, loss_head_m, HYDRAULIC_GRADE_CURVE_POINTS)
        except (TypeError, ValueError, ZeroDivisionError):
            return None

    def _embed_matplotlib_hydraulic_grade(
        self,
        doc,
        *,
        title,
        x_label,
        y_label,
        l_max,
        series_list,
        caption,
    ):
        """series_list: [(xs, ys, color, label), ...]；固定坐标轴范围与软件一致。"""
        try:
            plt, rcParams = _matplotlib_stack_for_export()
            rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial']
            rcParams['axes.unicode_minus'] = False

            all_y = []
            for xs, ys, _c, _lb in series_list:
                all_y.extend(ys)
            if not all_y:
                return False
            y_lo, y_hi = min(all_y), max(all_y)
            span = max(y_hi - y_lo, 1e-6)
            pad = max(span * 0.06, 0.5)
            y_min = y_lo - pad
            y_max = y_hi + pad

            fig, ax = plt.subplots(figsize=(7.2, 4.0), dpi=150)
            for xs, ys, color, label in series_list:
                ax.plot(xs, ys, color=color, linewidth=2.5, label=label)

            ax.set_xlim(0.0, float(l_max))
            ax.set_ylim(float(y_min), float(y_max))
            xticks = [l_max * i / HYDRAULIC_GRADE_TICK_DIVISIONS for i in range(HYDRAULIC_GRADE_TICK_DIVISIONS + 1)]
            ax.set_xticks(xticks)
            ax.set_xlabel(x_label, fontsize=11, fontstyle='italic')
            ax.set_ylabel(y_label, fontsize=11, fontstyle='italic')
            ax.set_title(title, fontsize=13, fontweight='bold', pad=10)
            ax.grid(True, linestyle='--', alpha=0.4)
            n_series = len(series_list)
            ncol = max(1, min(n_series, 3))
            ax.legend(
                fontsize=9,
                loc='upper center',
                bbox_to_anchor=(0.5, -0.14),
                ncol=ncol,
                frameon=True,
                columnspacing=1.4,
            )
            fig.subplots_adjust(left=0.11, right=0.97, top=0.90, bottom=0.22)

            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp_path = tmp.name
            tmp.close()
            fig.savefig(tmp_path, dpi=150)
            plt.close(fig)

            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(tmp_path, width=Inches(5.8))

            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                self._set_font(r)

            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            return True
        except ImportError:
            p = doc.add_paragraph(
                '（水力坡度图未能生成：安装包应已包含绘图组件；若开发环境运行请执行 pip install matplotlib，或重新安装应用。）'
            )
            for run in p.runs:
                self._set_font(run)
            return False

    def _add_slurry_total_head_process(self, doc, parameters, result):
        """浆体总扬程计算过程 + 与界面一致的双线水力坡度图（清水对比线由浆体参数推导）"""
        intermediate = result.get('intermediate', {})
        rho_k = parameters.get('rho_k', 'N/A')
        g = parameters.get('g', 9.81)
        H = parameters.get('H', 'N/A')
        rho_s = parameters.get('rho_s', 'N/A')
        i_k = parameters.get('i_k', 'N/A')
        L = parameters.get('L', 'N/A')
        P_j = parameters.get('P_j', 0)
        P_n = parameters.get('P_n', 0)
        P_z = parameters.get('P_z', 0)

        process_texts = [
            "公式: Pk = ρk·g·H + ρs·g·ik·L + Pj + Pn + Pz",
            f"1. 重力势能压力: ρk·g·H = {rho_k}×{g}×{H} = {intermediate.get('gravity_pressure', 'N/A')} kPa（ρ 单位 t/m³）",
            f"2. 沿程摩擦损失: ρs·g·ik·L = {rho_s}×{g}×{i_k}×{L} = {intermediate.get('friction_pressure', 'N/A')} kPa",
            f"3. 局部摩阻损失: Pj = {P_j} kPa",
            f"4. 泵站零件损失: Pn = {P_n} kPa",
            f"5. 出口余压: Pz = {P_z} kPa",
            f"6. 总输送压力: Pk = {intermediate.get('gravity_pressure', 0)} + {intermediate.get('friction_pressure', 0)} + {P_j} + {P_n} + {P_z} = {result.get('H_total', 'N/A')} kPa",
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

        doc.add_paragraph()
        hdg = doc.add_paragraph('附图：水力坡度线')
        for run in hdg.runs:
            run.bold = True
            self._set_font(run)

        slurry_xy = WordExporter._try_slurry_hydraulic_grade_xy(parameters)
        if not slurry_xy:
            p = doc.add_paragraph('（水力坡度图：当前参数不足以按界面模型绘制。）')
            for run in p.runs:
                self._set_font(run)
            return

        xs_s, ys_s = slurry_xy
        l_max = float(parameters.get('L') or 0)
        series = [(xs_s, ys_s, SLURRY_HYDRAULIC_LINE, '浆体水力坡度线')]
        cl_xy = WordExporter._try_slurry_page_clear_hydraulic_grade_xy(parameters)
        cap_extra = ''
        if cl_xy:
            series.append((cl_xy[0], cl_xy[1], CLEAR_HYDRAULIC_LINE, '清水对比水力坡度线'))
            cap_extra = (
                ' 橙色为浆体；蓝色为清水对比线（与浆体同 H、L、P_j、g，取 ρ_w=1 t/m³、i_w=i_k）。'
            )

        self._embed_matplotlib_hydraulic_grade(
            doc,
            title='浆体管道水力坡度线（示意）',
            x_label='管长 L (m)',
            y_label='水头 H (m)',
            l_max=l_max,
            series_list=series,
            caption=(
                '图：横轴管长 L，范围 [0, L_max]，主刻度步长 L_max/10；纵轴水头 H（m），范围按曲线极值加 6% 边距；Pn、Pz 未计入线内。'
                + cap_extra
            ),
        )

    def _add_clear_water_total_head_process(self, doc, parameters, result):
        """清水总扬程计算过程 + 与界面一致的单线水力坡度图"""
        intermediate = result.get('intermediate', {})
        rho_w = parameters.get('rho_w', 1)
        g = parameters.get('g', 9.81)
        H = parameters.get('H', 'N/A')
        i_w = parameters.get('i_w', 'N/A')
        L = parameters.get('L', 'N/A')
        P_j = parameters.get('P_j', 0)
        P_n = parameters.get('P_n', 0)
        P_z = parameters.get('P_z', 0)

        process_texts = [
            "公式: Pw = ρw·g·(H + i_w·L) + Pj + Pn + Pz",
            "注: 与分项 ρw·g·H + ρw·g·i_w·L 等价；ρk、ρs 在清水工况均取 ρw；沿程采用清水摩阻系数 i_w。",
            f"1. 重力势能压力: ρw·g·H = {rho_w}×{g}×{H} = {intermediate.get('gravity_pressure', 'N/A')} kPa（ρw 单位 t/m³）",
            f"2. 沿程摩擦损失: ρw·g·i_w·L = {rho_w}×{g}×{i_w}×{L} = {intermediate.get('friction_pressure', 'N/A')} kPa",
            f"3. 局部摩阻损失: Pj = {P_j} kPa",
            f"4. 泵站零件损失: Pn = {P_n} kPa",
            f"5. 出口余压: Pz = {P_z} kPa",
            f"6. 总输送压力: Pw = {intermediate.get('gravity_pressure', 0)} + {intermediate.get('friction_pressure', 0)} + {P_j} + {P_n} + {P_z} = {result.get('H_total', 'N/A')} kPa",
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

        doc.add_paragraph()
        hdg = doc.add_paragraph('附图：水力坡度线')
        for run in hdg.runs:
            run.bold = True
            self._set_font(run)

        l_max = float(parameters.get('L') or 0)
        cl_xy = WordExporter._try_clear_hydraulic_grade_xy(l_max, parameters)
        if not cl_xy:
            p = doc.add_paragraph('（水力坡度图：当前参数不足以按界面模型绘制。）')
            for run in p.runs:
                self._set_font(run)
            return

        xs, ys = cl_xy
        self._embed_matplotlib_hydraulic_grade(
            doc,
            title='清水管道水力坡度线（示意）',
            x_label='管长 L (m)',
            y_label='水头 H (m)',
            l_max=l_max,
            series_list=[(xs, ys, CLEAR_HYDRAULIC_LINE, '清水水力坡度线')],
            caption='图：横轴管长 L，范围 [0, L_max]，主刻度步长 L_max/10；纵轴水头 H（m），边距 6%；Pn、Pz 未计入线内。',
        )

    def _add_centrifugal_pump_total_head_process(self, doc, parameters, result):
        """离心泵总扬程：步骤1 K_p，步骤2 H_b = ΣH_s/(K_p·K_m)"""
        im = result.get('intermediate', {})
        cw = parameters.get('C_w', im.get('C_w', 'N/A'))
        kp = im.get('K_p', parameters.get('K_p', 'N/A'))
        s_hs = im.get('Sigma_H_s', parameters.get('Sigma_H_s', 'N/A'))
        km = im.get('K_m', parameters.get('K_m', 'N/A'))
        denom = im.get('K_p_K_m', 'N/A')
        hb = result.get('H_total', 'N/A')
        term = im.get('term_0p25_Cw')
        if term is None and cw not in (None, 'N/A'):
            try:
                cwf = float(cw)
                term = 0.25 * cwf
            except (TypeError, ValueError):
                term = 'N/A'
        process_texts = [
            "步骤1：主泵输送浆体时的扬程降低率 K_p = 1 - 0.25·C_w（C_w 为固相质量分数，浆体重量浓度）。",
            f"  已知 C_w = {cw}，0.25·C_w = {term if term is not None else 'N/A'}，得 K_p = {kp}。",
            "步骤2：主泵扬送清水的总扬程（液柱高度）H_b = ΣH_s / (K_p·K_m)。",
            f"  ΣH_s = {s_hs} m；K_m = {km}；K_p·K_m = {denom}。",
            f"  H_b = {s_hs} / ({denom}) = {hb} m。",
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)

    def _add_positive_displacement_pump_process(self, doc, parameters, result):
        """容积式泵：P_b = P_k / K_f"""
        im = result.get('intermediate', {})
        pk = im.get('P_k', parameters.get('P_k', 'N/A'))
        kf = im.get('K_f', parameters.get('K_f', 'N/A'))
        pb = result.get('P_b', result.get('H_total', 'N/A'))
        process_texts = [
            "容积式泵总扬程（压力形式）：P_b = P_k / K_f。",
            f"  已知 P_k = {pk} kPa，K_f = {kf}。",
            f"  P_b = {pk} / {kf} = {pb} kPa。",
        ]
        for text in process_texts:
            p = doc.add_paragraph(text)
            for run in p.runs:
                self._set_font(run)
        rho_k = parameters.get('rho_k', 1)
        g = parameters.get('g', 9.81)
        try:
            if pb != 'N/A' and rho_k and float(rho_k) > 0 and g and float(g) > 0:
                hm = float(pb) / (float(rho_k) * float(g))
                p2 = doc.add_paragraph(
                    f"折合浆体液柱高度（P_b/(ρ_k·g)）：约 {hm:.4f} m（ρ_k = {rho_k} t/m³，g = {g} m/s²）。"
                )
                for run in p2.runs:
                    self._set_font(run)
        except (TypeError, ValueError, ZeroDivisionError):
            pass

    def _add_software_promotion(self, doc):
        """附录：编制单位业务简介与致谢"""
        self._insert_page_break(doc)
        p = doc.add_paragraph()
        run = p.add_run("附录　编制单位简介")
        run.bold = True
        run.font.size = Pt(16)
        self._set_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        paras = [
            (
                "长沙有色冶金设计研究院有限公司（简称长沙有色院）成立于1953年，为国家高新技术企业、"
                "国家技术创新示范企业、国家企业技术中心，是我国成立较早的大型综合性设计研究单位之一，"
                "隶属于中国铝业集团有限公司，为中铝国际工程股份有限公司的子公司。"
            ),
            (
                "经过七十余年的发展，长沙有色院已形成覆盖有色金属行业全产业链与项目全生命周期的技术服务能力，"
                "持有冶金、市政、建筑、化工石化医药、环境工程及工程勘察、测绘、地质灾害治理等多类甲级资质，"
                "业务涵盖工程咨询、设计、工程总承包、监理、勘察、测绘、检验检测、施工、环境治理与生态修复、"
                "装备制造及科研开发等，在矿山、冶炼与环保等领域积累了大量工程经验与自主知识产权成果。"
            ),
            (
                "本院秉承「创新驱动，诚信服务，持续为客户创造价值」的理念，致力成为有色行业创新型领军企业。"
                "本浆体与清水管道水力计算软件由本院组织研发，用于设计辅助与资料整理，计算结论应与现行规范及工程实际相结合。"
            ),
            (
                "若将本计算书或本软件产出的结果用于工程设计依据、对外技术报审或结论性技术表述，须另行具备与长沙有色冶金设计研究院有限公司"
                "之间合法有效、与项目相符的正式业务合同，或该院出具的书面项目授权。本计算书与软件使用许可不替代上述合同或授权。"
            ),
        ]
        for text in paras:
            lp = doc.add_paragraph(text)
            for r in lp.runs:
                self._set_font(r)
            lp.paragraph_format.first_line_indent = Pt(24)
            lp.paragraph_format.line_spacing = 1.25

        doc.add_paragraph()
        thanks = doc.add_paragraph("感谢使用长沙有色冶金设计研究院有限公司浆体管道水力计算软件。")
        for r in thanks.runs:
            self._set_font(r)
        thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _get_unit(self, param_name):
        """根据参数名获取单位（与前端/API 常用字段对齐）"""
        units = {
            "D": "m",
            "D_n": "m",
            "ps": "t/m³",
            "pl": "t/m³",
            "ws": "m/s",
            "Cs": "decimal",
            "w0": "m/s",
            "phi": "",
            "FL": "",
            "K": "",
            "C": "",
            "Cg": "",
            "C1v": "",
            "theta": "度",
            "g": "m/s²",
            "G": "",
            "W": "",
            "rho_g": "t/m³",
            "rho_s": "t/m³",
            "rho_k": "t/m³",
            "rho_w": "t/m³",
            "dp": "mm",
            "beta": "",
            "lambda_coef": "",
            "V": "m/s",
            "C_w": "",
            "lambda_d": "",
            "L_s": "m",
            "L": "m",
            "H": "m",
            "i": "mH₂O/m",
            "i_w": "",
            "i_k": "mH₂O/m",
            "K_QL": "h²/m⁵",
            "Q": "m³/h",
            "P_j": "kPa",
            "P_n": "kPa",
            "P_z": "kPa",
            "Z1": "m",
            "Z2": "m",
            "H1": "m",
            "H2": "m",
            "d": "m",
            "eta_1": "Pa·s",
            "epsilon": "m",
        }
        return units.get(param_name, "")
